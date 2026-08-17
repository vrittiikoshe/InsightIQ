import fitz
from docx import Document as DocxDocument


def extract_pdf_text(file_path):
    text = ""

    try:
        doc = fitz.open(file_path)

        for page in doc:
            text += page.get_text()

        doc.close()

    except Exception as e:
        print("PDF Extraction Error:", e)

    return text


def extract_docx_text(file_path):
    text = ""

    try:
        doc = DocxDocument(file_path)

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Extract text from tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = []

                for cell in row.cells:
                    row_text.append(cell.text.strip())

                text += " | ".join(row_text) + "\n"

    except Exception as e:
        print("DOCX Extraction Error:", e)

    return text


def extract_txt_text(file_path):
    text = ""

    try:
        with open(
            file_path,
            "r",
            encoding="utf-8",
            errors="ignore"
        ) as file:
            text = file.read()

    except Exception as e:
        print("TXT Extraction Error:", e)

    return text


def extract_document_text(file_path, file_type):
    """
    Extract text based on document type.
    """

    if file_type == "PDF":
        return extract_pdf_text(file_path)

    elif file_type == "DOCX":
        return extract_docx_text(file_path)

    elif file_type == "TXT":
        return extract_txt_text(file_path)

    return ""